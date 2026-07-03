"""Tests for core.word_exporter."""
import json
import os
import pytest
from datetime import date
from pathlib import Path
from core.word_exporter import (
    create_invoice_report,
    _fetch_boi_rate,
    _find_business_day,
    _get_exchange_rates,
)


@pytest.fixture
def sample_rows():
    return [
        {"date": "2024-01-15", "description": "Hostinger VPS", "amount": 89.00, "currency": "₪", "notes": "חודשי"},
        {"date": "2024-01-22", "description": "Google Cloud", "amount": 142.50, "currency": "₪", "notes": ""},
        {"date": "2024-02-03", "description": "AWS Services", "amount": 67.30, "currency": "₪", "notes": "ניסיון"},
    ]


@pytest.fixture
def tmp_exports(tmp_path):
    return tmp_path / "exports"


class TestCreateInvoiceReport:
    def test_creates_docx_file(self, sample_rows, tmp_exports):
        path = create_invoice_report(sample_rows, output_dir=str(tmp_exports))
        assert path.endswith(".docx")
        assert os.path.isfile(path)

    def test_file_not_empty(self, sample_rows, tmp_exports):
        path = create_invoice_report(sample_rows, output_dir=str(tmp_exports))
        assert os.path.getsize(path) > 0

    def test_contains_correct_row_count(self, sample_rows, tmp_exports):
        from docx import Document
        path = create_invoice_report(sample_rows, output_dir=str(tmp_exports))
        doc = Document(path)
        table = doc.tables[0]
        # header + 3 data rows + 1 summary = 5
        assert len(table.rows) == 5

    def test_summary_row_total(self, sample_rows, tmp_exports):
        from docx import Document
        path = create_invoice_report(sample_rows, output_dir=str(tmp_exports))
        doc = Document(path)
        table = doc.tables[0]
        last_row = table.rows[-1]
        # The report table is 3 columns (date, amount, supplier); the total
        # lives in the amount cell. Check the whole summary row so the test is
        # robust to column-layout changes rather than pinning a fixed index.
        row_text = " ".join(c.text for c in last_row.cells)
        assert "298.80" in row_text

    def test_empty_rows_returns_none(self, tmp_exports):
        result = create_invoice_report([], output_dir=str(tmp_exports))
        assert result is None

    def test_creates_output_dir(self, sample_rows, tmp_path):
        new_dir = str(tmp_path / "new_dir" / "exports")
        path = create_invoice_report(sample_rows, output_dir=new_dir)
        assert os.path.isfile(path)


# ── M21: Israeli business-day walk-back + BOI window hardening ──────────────
# 2026-06-25=Thu, 06-26=Fri, 06-27=Sat, 06-28=Sun, 06-29=Mon.


class TestFindBusinessDay:
    """Bank of Israel publishes no rates on Fri/Sat (the Israeli weekend).
    The old walk-back skipped Sat/Sun: a Friday invoice queried Friday
    (no published rate -> 'not available'), and Sunday — a business day in
    Israel — was wrongly walked back."""

    def test_friday_walks_back_to_thursday(self):
        d, was_weekend = _find_business_day(date(2026, 6, 26))
        assert d == date(2026, 6, 25)
        assert was_weekend is True

    def test_saturday_walks_back_to_thursday(self):
        d, was_weekend = _find_business_day(date(2026, 6, 27))
        assert d == date(2026, 6, 25)
        assert was_weekend is True

    def test_sunday_is_a_business_day(self):
        d, was_weekend = _find_business_day(date(2026, 6, 28))
        assert d == date(2026, 6, 28)
        assert was_weekend is False

    def test_monday_unchanged(self):
        d, was_weekend = _find_business_day(date(2026, 6, 29))
        assert d == date(2026, 6, 29)
        assert was_weekend is False


class _FakeResponse:
    def __init__(self, payload: dict):
        self._payload = payload

    def read(self):
        return json.dumps(self._payload).encode()

    def __enter__(self):
        return self

    def __exit__(self, *args):
        return False


def _boi_payload(dates_rates: list[tuple[str, float]]) -> dict:
    """Minimal BOI SDMX-JSON shape: observations keyed by index, with the
    matching TIME_PERIOD values in the structure block."""
    return {
        "data": {
            "dataSets": [{"series": {"0:0:0:0": {"observations": {
                str(i): [str(rate)] for i, (_, rate) in enumerate(dates_rates)
            }}}}],
            "structures": [{"dimensions": {"observation": [{
                "id": "TIME_PERIOD",
                "values": [{"id": ds, "start": ds} for ds, _ in dates_rates],
            }]}}],
        }
    }


class TestFetchBoiRate:
    def test_windowed_query_returns_latest_observation_and_its_date(self, monkeypatch):
        captured = {}
        payload = _boi_payload([("2026-06-24", 3.61), ("2026-06-25", 3.64)])

        def fake_urlopen(req, timeout=10):
            captured["url"] = req.full_url
            return _FakeResponse(payload)

        monkeypatch.setattr("core.word_exporter.urllib.request.urlopen", fake_urlopen)
        rate, rate_date = _fetch_boi_rate(date(2026, 6, 26))  # Friday
        assert rate == 3.64
        assert rate_date == date(2026, 6, 25)
        # Trailing 7-day window ending at the nearest business day (Thursday)
        assert "startperiod=2026-06-18" in captured["url"]
        assert "endperiod=2026-06-25" in captured["url"]

    def test_empty_observations_returns_none_without_raising(self, monkeypatch):
        payload = {"data": {
            "dataSets": [{"series": {"0:0:0:0": {"observations": {}}}}],
            "structures": [],
        }}
        monkeypatch.setattr(
            "core.word_exporter.urllib.request.urlopen",
            lambda req, timeout=10: _FakeResponse(payload),
        )
        assert _fetch_boi_rate(date(2026, 6, 26)) == (None, None)

    def test_network_error_returns_none_pair(self, monkeypatch):
        def boom(req, timeout=10):
            raise OSError("no network")

        monkeypatch.setattr("core.word_exporter.urllib.request.urlopen", boom)
        assert _fetch_boi_rate(date(2026, 6, 26)) == (None, None)

    def test_missing_structure_falls_back_to_business_day(self, monkeypatch):
        payload = _boi_payload([("2026-06-25", 3.64)])
        del payload["data"]["structures"]  # no way to resolve the obs date
        monkeypatch.setattr(
            "core.word_exporter.urllib.request.urlopen",
            lambda req, timeout=10: _FakeResponse(payload),
        )
        rate, rate_date = _fetch_boi_rate(date(2026, 6, 26))  # Friday
        assert rate == 3.64
        assert rate_date == date(2026, 6, 25)  # walk-back Thursday


class TestGetExchangeRates:
    def test_friday_usd_invoice_reports_actual_rate_day(self, monkeypatch):
        payload = _boi_payload([("2026-06-25", 3.64)])
        monkeypatch.setattr(
            "core.word_exporter.urllib.request.urlopen",
            lambda req, timeout=10: _FakeResponse(payload),
        )
        rows = [{"date": "2026-06-26", "amount": 10.0, "currency": "USD"}]
        assert _get_exchange_rates(rows) == [{
            "date": "2026-06-26",
            "biz_day": "2026-06-25",
            "was_weekend": True,
            "rate": 3.64,
        }]

    def test_fetch_failure_still_reports_walkback_day(self, monkeypatch):
        def boom(req, timeout=10):
            raise OSError("down")

        monkeypatch.setattr("core.word_exporter.urllib.request.urlopen", boom)
        rows = [{"date": "2026-06-27", "amount": 5.0, "currency": "USD"}]  # Saturday
        entry = _get_exchange_rates(rows)[0]
        assert entry["rate"] is None
        assert entry["biz_day"] == "2026-06-25"
        assert entry["was_weekend"] is True

    def test_non_usd_rows_do_not_fetch(self, monkeypatch):
        def boom(req, timeout=10):
            raise AssertionError("BOI must not be queried for non-USD rows")

        monkeypatch.setattr("core.word_exporter.urllib.request.urlopen", boom)
        assert _get_exchange_rates([{"date": "2026-06-26", "currency": "₪"}]) == []
