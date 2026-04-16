"""
Export invoices to a professional Word document (.docx).

Structure:
  1. Title: "דוח קבלות" (bold, large, RTL)
  2. Subtitle: date range of invoices
  3. Main table: תאריך / סכום / ספק + total row
  4. Notes section: receipt counts by currency
  5. Exchange rates section: BOI rates for USD invoices
  6. RTL Hebrew throughout, professional formatting
"""

import os
import logging
import urllib.request
import json
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

HEBREW_MONTHS = {
    1: "\u05d9\u05e0\u05d5\u05d0\u05e8",      # ינואר
    2: "\u05e4\u05d1\u05e8\u05d5\u05d0\u05e8",  # פברואר
    3: "\u05de\u05e8\u05e5",                      # מרץ
    4: "\u05d0\u05e4\u05e8\u05d9\u05dc",          # אפריל
    5: "\u05de\u05d0\u05d9",                      # מאי
    6: "\u05d9\u05d5\u05e0\u05d9",                # יוני
    7: "\u05d9\u05d5\u05dc\u05d9",                # יולי
    8: "\u05d0\u05d5\u05d2\u05d5\u05e1\u05d8",    # אוגוסט
    9: "\u05e1\u05e4\u05d8\u05de\u05d1\u05e8",    # ספטמבר
    10: "\u05d0\u05d5\u05e7\u05d8\u05d5\u05d1\u05e8",  # אוקטובר
    11: "\u05e0\u05d5\u05d1\u05de\u05d1\u05e8",  # נובמבר
    12: "\u05d3\u05e6\u05de\u05d1\u05e8",          # דצמבר
}


# ── Currency symbol → ISO code mapping ──────────────────────────────────────
_SYMBOL_TO_ISO = {"₪": "ILS", "$": "USD", "€": "EUR", "£": "GBP"}

def _norm_currency(raw: str | None) -> str:
    """Normalize currency symbol or code to uppercase ISO 4217 code."""
    code = (raw or "ILS").upper()
    return _SYMBOL_TO_ISO.get(code, code)


# ── RTL helpers ──────────────────────────────────────────────────────────────

def _set_rtl(paragraph):
    """Set paragraph direction to RTL."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)


def _set_cell_rtl(cell):
    """Set all paragraphs in a table cell to RTL."""
    for p in cell.paragraphs:
        _set_rtl(p)


def _set_table_rtl(table):
    """Set table visual layout to RTL (columns flow right-to-left)."""
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = table._tbl._add_tblPr()
    bidi = tbl_pr.makeelement(qn("w:bidiVisual"), {})
    tbl_pr.append(bidi)


def _set_section_rtl(section):
    """Set the document section to RTL."""
    sect_pr = section._sectPr
    bidi = sect_pr.makeelement(qn("w:bidi"), {})
    sect_pr.append(bidi)


# ── Currency formatting ─────────────────────────────────────────────────────

def _format_amount(amount: float | None, currency: str = "ILS") -> str:
    if amount is None:
        return "\u2014"
    symbols = {"ILS": "\u20aa", "USD": "$", "EUR": "\u20ac", "GBP": "\u00a3"}
    sym = symbols.get(currency, currency + " ")
    return f"{sym}{amount:,.2f}"


# ── Date helpers ─────────────────────────────────────────────────────────────

def _parse_date(d) -> date | None:
    if isinstance(d, date):
        return d
    if isinstance(d, datetime):
        return d.date()
    if isinstance(d, str):
        for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%d/%m/%Y"):
            try:
                return datetime.strptime(d[:10], fmt).date()
            except ValueError:
                continue
    return None


def _date_range_subtitle(rows: list[dict]) -> str:
    """Build Hebrew date range string like 'ינואר 2026 – אפריל 2026'."""
    dates = [_parse_date(r.get("date")) for r in rows]
    dates = [d for d in dates if d is not None]
    if not dates:
        return ""
    min_d, max_d = min(dates), max(dates)
    start = f"{HEBREW_MONTHS[min_d.month]} {min_d.year}"
    end = f"{HEBREW_MONTHS[max_d.month]} {max_d.year}"
    if start == end:
        return start
    return f"{start} \u2013 {end}"


# ── Bank of Israel exchange rates ────────────────────────────────────────────

def _find_business_day(d: date) -> tuple[date, bool]:
    """Walk back to the nearest business day (Mon-Fri). Returns (date, was_weekend)."""
    original = d
    while d.weekday() >= 5:  # 5=Sat, 6=Sun
        d -= timedelta(days=1)
    return d, d != original


def _fetch_boi_rate(d: date) -> float | None:
    """Fetch USD/ILS exchange rate from Bank of Israel for a given date."""
    biz_day, _ = _find_business_day(d)
    ds = biz_day.strftime("%Y-%m-%d")
    url = (
        f"https://edge.boi.gov.il/FusionEdgeServer/sdmx/v2/data/dataflow/"
        f"BOI.STATISTICS/EXR/1.0/RER_USD_ILS"
        f"?startperiod={ds}&endperiod={ds}&format=sdmx-json"
    )
    try:
        req = urllib.request.Request(url, headers={"Accept": "application/json"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read())
        observations = data["data"]["dataSets"][0]["series"]["0:0:0:0"]["observations"]
        for key in observations:
            return float(observations[key][0])
    except Exception as e:
        logger.warning("Failed to fetch BOI rate for %s: %s", ds, e)
    return None


def _get_exchange_rates(rows: list[dict]) -> list[dict]:
    """For each USD invoice, fetch the BOI exchange rate. Returns list of rate info dicts."""
    usd_dates: dict[str, date] = {}
    for r in rows:
        if _norm_currency(r.get("currency")) != "USD":
            continue
        d = _parse_date(r.get("date"))
        if d:
            ds = d.strftime("%Y-%m-%d")
            if ds not in usd_dates:
                usd_dates[ds] = d

    rates = []
    for ds, d in sorted(usd_dates.items()):
        biz_day, was_weekend = _find_business_day(d)
        rate = _fetch_boi_rate(d)
        entry = {
            "date": ds,
            "biz_day": biz_day.strftime("%Y-%m-%d"),
            "was_weekend": was_weekend,
            "rate": rate,
        }
        rates.append(entry)
    return rates


# ── Table styling helpers ────────────────────────────────────────────────────

def _set_cell_borders(cell, color="4472C4"):
    """Set thin borders on a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    for edge in ("top", "left", "bottom", "right"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): color,
        })
        borders.append(el)
    tc_pr.append(borders)


def _shade_cell(cell, color: str):
    """Set cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    tc_pr.append(shading)


# ══════════════════════════════════════════════════════════════════════════════
# MAIN EXPORT FUNCTION
# ══════════════════════════════════════════════════════════════════════════════

def create_invoice_report(
    rows: list[dict],
    output_dir: str = "exports",
    filename: str | None = None,
    organization_name: str | None = None,
) -> str | None:
    """Create a professional .docx invoice report in Hebrew RTL format.

    Args:
        rows: list of invoice dicts with keys:
            id, company, subject, sender, amount, currency,
            date, classification_tier, has_attachment, scan_id, notes
        output_dir: directory to save the file
        filename: optional filename override
        organization_name: org name for the page header

    Returns:
        Path to created .docx file, or None if rows is empty.
    """
    if not rows:
        return None

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    if filename is None:
        filename = f"invoices_report_{date.today()}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()

    # -- Global styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # -- RTL section --
    section = doc.sections[0]
    _set_section_rtl(section)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # -- Page header --
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    _set_rtl(header_para)
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    org_label = organization_name or ""
    header_run = header_para.add_run(
        f"{org_label}  \u2022  {datetime.now().strftime('%d/%m/%Y')}" if org_label
        else datetime.now().strftime('%d/%m/%Y')
    )
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── 1. Title ─────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    _set_rtl(title_para)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    title_run = title_para.add_run("\u05d3\u05d5\u05d7 \u05e7\u05d1\u05dc\u05d5\u05ea")  # דוח קבלות
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # ── 2. Subtitle — date range ─────────────────────────────────────────
    subtitle_text = _date_range_subtitle(rows)
    if subtitle_text:
        sub_para = doc.add_paragraph()
        _set_rtl(sub_para)
        sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        sub_run = sub_para.add_run(subtitle_text)
        sub_run.font.size = Pt(13)
        sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph("")  # spacer

    # ── 3. Main table — תאריך / סכום / ספק ───────────────────────────────
    # Headers (RTL order: rightmost column first visually)
    col_headers = [
        "\u05ea\u05d0\u05e8\u05d9\u05da",  # תאריך
        "\u05e1\u05db\u05d5\u05dd",          # סכום
        "\u05e1\u05e4\u05e7",                # ספק
    ]

    table = doc.add_table(rows=1, cols=3)
    _set_table_rtl(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(col_headers):
        cell = hdr_cells[i]
        cell.text = h
        _set_cell_rtl(cell)
        _set_cell_borders(cell, "2F5496")
        _shade_cell(cell, "2F5496")
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Compute totals for ILS
    total_ils = 0.0
    ils_count = 0
    usd_count = 0
    for r in rows:
        currency = _norm_currency(r.get("currency"))
        amount = r.get("amount") or 0
        if currency == "ILS":
            total_ils += amount
            ils_count += 1
        elif currency == "USD":
            usd_count += 1

    # Data rows
    for idx, row in enumerate(rows):
        d = _parse_date(row.get("date"))
        date_str = d.strftime("%d/%m/%Y") if d else "\u2014"
        amount = row.get("amount")
        currency = _norm_currency(row.get("currency"))
        amount_str = _format_amount(amount, currency)
        supplier = row.get("company") or row.get("description") or "\u2014"

        cells = table.add_row().cells
        cell_values = [date_str, amount_str, supplier]
        stripe_color = "F2F6FC" if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate(cell_values):
            cell = cells[i]
            cell.text = val
            _set_cell_rtl(cell)
            _set_cell_borders(cell, "B4C6E7")
            _shade_cell(cell, stripe_color)
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    # Total row
    total_cells = table.add_row().cells
    total_values = [
        "\u05e1\u05d4\u05f4\u05db",  # סה״כ
        _format_amount(total_ils, "ILS"),
        "",
    ]
    for i, val in enumerate(total_values):
        cell = total_cells[i]
        cell.text = val
        _set_cell_rtl(cell)
        _set_cell_borders(cell, "2F5496")
        _shade_cell(cell, "D6E4F0")
        for p in cell.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    doc.add_paragraph("")  # spacer

    # ── 4. Notes section — הערות ──────────────────────────────────────────
    notes_heading = doc.add_paragraph()
    _set_rtl(notes_heading)
    notes_heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = notes_heading.add_run("\u05d4\u05e2\u05e8\u05d5\u05ea")  # הערות
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    notes_lines = [
        f"\u05e1\u05d4\u05f4\u05db \u05e7\u05d1\u05dc\u05d5\u05ea \u05d1\u05d3\u05d5\u05d7: {len(rows)}",  # סה״כ קבלות בדוח: X
        f'\u05e7\u05d1\u05dc\u05d5\u05ea \u05d1\u05e9"\u05d7: {ils_count}',  # קבלות בש"ח: X
        f"\u05e7\u05d1\u05dc\u05d5\u05ea \u05d1\u05d3\u05d5\u05dc\u05e8: {usd_count}",  # קבלות בדולר: X
    ]
    for line in notes_lines:
        p = doc.add_paragraph()
        _set_rtl(p)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r = p.add_run(line)
        r.font.size = Pt(10)

    doc.add_paragraph("")  # spacer

    # ── 5. Exchange rates section — שערי חליפין ───────────────────────────
    if usd_count > 0:
        rates_heading = doc.add_paragraph()
        _set_rtl(rates_heading)
        rates_heading.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = rates_heading.add_run("\u05e9\u05e2\u05e8\u05d9 \u05d7\u05dc\u05d9\u05e4\u05d9\u05df")  # שערי חליפין
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        exchange_rates = _get_exchange_rates(rows)
        if exchange_rates:
            rate_table = doc.add_table(rows=1, cols=3)
            _set_table_rtl(rate_table)
            rate_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            rate_headers = [
                "\u05ea\u05d0\u05e8\u05d9\u05da",  # תאריך
                "\u05e9\u05e2\u05e8 \u05d3\u05d5\u05dc\u05e8/\u05e9\u05e7\u05dc",  # שער דולר/שקל
                "\u05d4\u05e2\u05e8\u05d5\u05ea",  # הערות
            ]
            for i, h in enumerate(rate_headers):
                cell = rate_table.rows[0].cells[i]
                cell.text = h
                _set_cell_rtl(cell)
                _set_cell_borders(cell, "2F5496")
                _shade_cell(cell, "2F5496")
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for entry in exchange_rates:
                cells = rate_table.add_row().cells
                cells[0].text = entry["date"]
                cells[1].text = f"{entry['rate']:.4f}" if entry["rate"] else "\u05dc\u05d0 \u05d6\u05de\u05d9\u05df"  # לא זמין
                if entry["was_weekend"]:
                    cells[2].text = f"\u05e1\u05d5\u05e3 \u05e9\u05d1\u05d5\u05e2 \u2014 \u05e0\u05dc\u05e7\u05d7 \u05e9\u05e2\u05e8 \u05de\u05d9\u05d5\u05dd {entry['biz_day']}"  # סוף שבוע — נלקח שער מיום
                else:
                    cells[2].text = ""
                for cell in cells:
                    _set_cell_rtl(cell)
                    _set_cell_borders(cell, "B4C6E7")
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in p.runs:
                            run.font.size = Pt(9)
        else:
            p = doc.add_paragraph()
            _set_rtl(p)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            r = p.add_run("\u05dc\u05d0 \u05e0\u05d9\u05ea\u05df \u05dc\u05d0\u05d7\u05d6\u05e8 \u05e9\u05e2\u05e8\u05d9 \u05d7\u05dc\u05d9\u05e4\u05d9\u05df \u05de\u05d1\u05e0\u05e7 \u05d9\u05e9\u05e8\u05d0\u05dc")  # לא ניתן לאחזר שערי חליפין מבנק ישראל
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Footer ────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer_para.add_run(
        f"\u05d4\u05d5\u05e4\u05e7 \u05e2\u05dc \u05d9\u05d3\u05d9 Invoice Fetcher  \u2022  "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.save(filepath)
    return filepath
